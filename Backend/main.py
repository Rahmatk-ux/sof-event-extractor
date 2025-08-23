from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
import uvicorn, os, io, re, datetime as dt

# PDF & DOCX
import fitz  # PyMuPDF
from docx import Document
import pandas as pd

app = FastAPI(title="SoF Event Extractor")

# --- Allow your React app to call the API during dev ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Utility: read file into temp path ---
def save_temp(upload: UploadFile) -> str:
    name = f"temp_{upload.filename}"
    with open(name, "wb") as f:
        f.write(upload.file.read())
    return name

# --- Read text from PDF ---
def read_pdf(path: str) -> str:
    text = []
    with fitz.open(path) as doc:
        for page in doc:
            text.append(page.get_text("text"))
    return "\n".join(text)

# --- Read text from DOCX ---
def read_docx(path: str) -> str:
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                lines.append(row_text)
    return "\n".join(lines)

# --- Heuristics ---
EVENT_WORDS = [
    "loading","discharging","anchorage","berthing","unberthing","shifting",
    "arrival","arrived","departure","departed","sailing","bunkering",
    "weather","nor","notice of readiness","draft survey","customs","immigration",
    "hatch","stop","resume","suspension","cargo operations","ballast"
]

TIME_RE = re.compile(r"\b([01]?\d|2[0-3]):[0-5]\d(?::[0-5]\d)?\b")
DATE_RE = re.compile(r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{1,2}\s+[A-Za-z]{3,}\s+\d{2,4})\b", re.I)

def normalize_date(s: str) -> str | None:
    s = s.strip()
    for fmt in ("%d/%m/%Y","%d-%m-%Y","%d/%m/%y","%d-%m-%y","%d %b %Y","%d %B %Y"):
        try:
            return dt.datetime.strptime(s, fmt).date().isoformat()
        except Exception:
            pass
    return None

def line_has_event(line: str) -> str | None:
    low = line.lower()
    for w in EVENT_WORDS:
        if w in low:
            return w
    return None

def extract_events(text: str):
    events = []
    current_date_iso = None
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        md = DATE_RE.search(line)
        if md:
            iso = normalize_date(md.group(0))
            if iso:
                current_date_iso = iso

        times = TIME_RE.findall(line)
        ev_word = line_has_event(line)
        if not ev_word and not times:
            continue

        if ev_word and len(times) >= 2:
            start_t = times[0]
            end_t = times[1]
            events.append({
                "event": ev_word.title(),
                "start": f"{current_date_iso} {start_t}" if current_date_iso else start_t,
                "end": f"{current_date_iso} {end_t}" if current_date_iso else end_t,
                "source": line
            })
        elif ev_word and len(times) == 1:
            events.append({
                "event": ev_word.title(),
                "start": f"{current_date_iso} {times[0]}" if current_date_iso else times[0],
                "end": None,
                "source": line
            })
        elif not ev_word and len(times) >= 2:
            events.append({
                "event": "Timed Activity",
                "start": f"{current_date_iso} {times[0]}" if current_date_iso else times[0],
                "end": f"{current_date_iso} {times[1]}" if current_date_iso else times[1],
                "source": line
            })
    return events

# --- API Endpoints ---
@app.post("/extract")
async def extract(file: UploadFile = File(...)):
    if not (file.filename.lower().endswith(".pdf") or file.filename.lower().endswith(".docx")):
        raise HTTPException(400, detail="Upload a .pdf or .docx")

    path = save_temp(file)
    try:
        text = read_pdf(path) if path.lower().endswith(".pdf") else read_docx(path)
        events = extract_events(text)
        return {"count": len(events), "events": events}
    finally:
        try: os.remove(path)
        except: pass

@app.post("/extract/csv")
async def extract_csv(file: UploadFile = File(...)):
    path = save_temp(file)
    try:
        text = read_pdf(path) if path.lower().endswith(".pdf") else read_docx(path)
        rows = extract_events(text)
        df = pd.DataFrame(rows)
        buf = io.StringIO()
        df.to_csv(buf, index=False)
        buf.seek(0)
        return StreamingResponse(iter([buf.getvalue()]),
                                 media_type="text/csv",
                                 headers={"Content-Disposition":"attachment; filename=events.csv"})
    finally:
        try: os.remove(path)
        except: pass

# --- Serve Frontend ---
frontend_path = os.path.join(os.path.dirname(__file__), "frontend")
app.mount("/assets", StaticFiles(directory=os.path.join(frontend_path, "assets")), name="assets")

@app.get("/", response_class=HTMLResponse)
async def serve_index():
    with open(os.path.join(frontend_path, "index.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

@app.get("/{full_path:path}", response_class=HTMLResponse)
async def catch_all(full_path: str):
    # React router fallback
    with open(os.path.join(frontend_path, "index.html"), "r", encoding="utf-8") as f:
        return HTMLResponse(content=f.read())

if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8000)
